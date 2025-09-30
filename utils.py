import os
import re
import json
import logging
import tempfile
from pathlib import Path
from typing import Dict, List, Tuple, Optional
from docx import Document
from pathlib import Path
import json
from typing import Dict, List, Tuple, Optional, Any, Set
import asyncio
import time


try:
    from config import (
        OWNER_ID, ADMIN_USER_IDS, bot,
        DATA_DIR, TESTS_DIR, STUDENTS_DIR,
        ACTIVE_TEST_FILE, GROUPS_FILE, STUDENTS_FILE,
        USER_GROUPS_FILE, GROUP_MEMBERS_FILE,
    )
except Exception:
    OWNER_ID = int(os.environ.get("OWNER_ID", "0")) or None
    ADMIN_USER_IDS = [OWNER_ID] if OWNER_ID else []
    DATA_DIR = "data"
    TESTS_DIR = os.path.join(DATA_DIR, "tests")
    STUDENTS_DIR = os.path.join(DATA_DIR, "students")
    ACTIVE_TEST_FILE = os.path.join(TESTS_DIR, "active_tests.json")
    GROUPS_FILE = os.path.join(DATA_DIR, "groups.txt")
    STUDENTS_FILE = os.path.join(DATA_DIR, "students.json")
    USER_GROUPS_FILE = os.path.join(DATA_DIR, "user_groups.json")
    GROUP_MEMBERS_FILE = os.path.join(DATA_DIR, "group_members.json")

log = logging.getLogger("utils")
if not log.handlers:
    h = logging.StreamHandler()
    h.setFormatter(logging.Formatter("%(asctime)s %(levelname)s %(name)s: %(message)s"))
    log.addHandler(h)
log.setLevel(logging.INFO)

def ensure_dir(p: Path):
    p.mkdir(parents=True, exist_ok=True)

def read_json(path: Path, default):
    try:
        if path.exists():
            return json.loads(path.read_text(encoding="utf-8"))
    except Exception as e:
        log.error(f"read_json({path}): {e}")
    return default

def write_json(path: Path, obj):
    try:
        ensure_dir(path.parent)
        path.write_text(json.dumps(obj, ensure_ascii=False, indent=2), encoding="utf-8")
        return True
    except Exception as e:
        log.error(f"write_json({path}): {e}")
        return False

def get_file_path(base_dir: str, filename: str) -> str:
    p = Path(base_dir) / filename
    ensure_dir(p.parent)
    return str(p)

def ensure_data():
    ensure_dir(Path(DATA_DIR))
    ensure_dir(Path(TESTS_DIR))
    ensure_dir(Path(STUDENTS_DIR))
    if not Path(ACTIVE_TEST_FILE).exists():
        write_json(Path(ACTIVE_TEST_FILE), {"active_tests": []})
    for f, default in [
        (STUDENTS_FILE, {}),
        (USER_GROUPS_FILE, {}),
        (GROUP_MEMBERS_FILE, {}),
    ]:
        p = Path(f)
        if not p.exists():
            write_json(p, default)
    Path(GROUPS_FILE).parent.mkdir(parents=True, exist_ok=True)
    if not Path(GROUPS_FILE).exists():
        Path(GROUPS_FILE).write_text("", encoding="utf-8")

def is_owner(uid: int) -> bool:
    if OWNER_ID and uid == int(OWNER_ID):
        return True
    return int(uid) in set(int(x) for x in (ADMIN_USER_IDS or []))

def _parse_groups_file(text: str) -> Dict[int, str]:
    out: Dict[int, str] = {}
    for line in text.splitlines():
        line = line.strip()
        if not line:
            continue
        if "," in line:
            id_str, title = line.split(",", 1)
        else:
            id_str, title = line, ""
        try:
            out[int(id_str.strip())] = title.strip()
        except Exception:
            continue
    return out

def load_group_ids() -> List[int]:
    if not Path(GROUPS_FILE).exists():
        return []
    try:
        data = Path(GROUPS_FILE).read_text(encoding="utf-8")
    except Exception:
        return []
    return list(_parse_groups_file(data).keys())

def load_group_titles() -> Dict[int, str]:
    if not Path(GROUPS_FILE).exists():
        return {}
    data = Path(GROUPS_FILE).read_text(encoding="utf-8")
    return _parse_groups_file(data)

def add_or_update_group(chat_id: int, title: str = ""):
    groups = load_group_titles()
    groups[int(chat_id)] = (title or groups.get(int(chat_id), "")).strip()
    lines = []
    for gid, t in sorted(groups.items()):
        lines.append(f"{gid},{t}" if t else f"{gid}")
    Path(GROUPS_FILE).write_text("\n".join(lines), encoding="utf-8")

def remove_group(chat_id: int):
    groups = load_group_titles()
    if int(chat_id) in groups:
        del groups[int(chat_id)]
        lines = []
        for gid, t in sorted(groups.items()):
            lines.append(f"{gid},{t}" if t else f"{gid}")
        Path(GROUPS_FILE).write_text("\n".join(lines), encoding="utf-8")

def load_group_members() -> Dict[str, Dict[str, List[int]]]:
    return read_json(Path(GROUP_MEMBERS_FILE), {})

def get_group_member_ids(group_id: int) -> List[int]:
    gm = load_group_members()
    rec = gm.get(str(group_id)) or {}
    return [int(x) for x in rec.get("members", [])]

def update_group_member(group_id: int, user_id: int, present: bool):
    gm = load_group_members()
    rec = gm.get(str(group_id)) or {"members": []}
    s = set(int(x) for x in rec.get("members", []))
    if present:
        s.add(int(user_id))
    else:
        s.discard(int(user_id))
    rec["members"] = sorted(s)
    gm[str(group_id)] = rec
    write_json(Path(GROUP_MEMBERS_FILE), gm)

async def sync_group_members(group_id: int) -> int:
    """Sync members using user account if available, else bot account"""
    try:
        from telethon_service import get_user_telethon_service, get_telethon_service

        # 1) Try user Telethon (can usually fetch all members)
        try:
            user_telethon = await get_user_telethon_service()
        except Exception:
            user_telethon = None

        if user_telethon:
            try:
                members, total_count = await user_telethon.fetch_all_group_members(group_id)
            except Exception as e:
                log.warning(f"User Telethon fetch_all_group_members failed for {group_id}: {e}")
                members = None
                total_count = 0

            if members:
                member_ids = []
                member_data: Dict[int, Dict] = {}
                for member in members:
                    # skip bots if reported
                    if member.get("is_bot"):
                        continue
                    uid = int(member["id"])
                    member_ids.append(uid)
                    member_data[uid] = {
                        "first_name": member.get("first_name", "") or "",
                        "last_name": member.get("last_name", "") or "",
                        "username": member.get("username", "") or "",
                        "phone": member.get("phone", "") or "",
                        "is_premium": member.get("is_premium", False),
                        "is_bot": member.get("is_bot", False),
                        "is_admin": member.get("is_admin", False),
                    }

                # Load existing data to preserve information
                gm = load_group_members()
                existing_data = gm.get(str(group_id), {}).get("member_data", {})
                
                # Merge data: preserve existing info for users who might have incomplete new data
                for uid_str, existing_info in existing_data.items():
                    uid = int(uid_str)
                    if uid in member_data:
                        # Preserve non-empty fields from existing data
                        for field in ["phone", "first_name", "last_name", "username"]:
                            if not member_data[uid].get(field) and existing_info.get(field):
                                member_data[uid][field] = existing_info[field]

                import time
                gm[str(group_id)] = {
                    "members": sorted(member_ids),
                    "member_data": {str(k): v for k, v in member_data.items()},
                    "last_sync": int(time.time()),
                    "sync_method": "user_telethon",
                    "total_count": total_count or len(member_ids),
                }
                write_json(Path(GROUP_MEMBERS_FILE), gm)

                # update user_groups.json
                user_groups_data = read_json(Path(USER_GROUPS_FILE), {})
                for user_id in member_ids:
                    uid_s = str(user_id)
                    if uid_s not in user_groups_data:
                        user_groups_data[uid_s] = []
                    if group_id not in user_groups_data[uid_s]:
                        user_groups_data[uid_s].append(group_id)
                write_json(Path(USER_GROUPS_FILE), user_groups_data)

                log.info(f"User Telethon synced {len(member_ids)} members for group {group_id}")
                return len(member_ids)

        # 2) Try bot Telethon client
        try:
            telethon = await get_telethon_service()
        except Exception:
            telethon = None

        if telethon:
            try:
                # Use fetch_group_members (may return subset depending on bot permissions)
                members, total_count = await telethon.fetch_group_members(
                    group_id, exclude_admins=False, exclude_bots=True
                )
            except Exception as e:
                log.warning(f"Bot Telethon fetch_group_members failed for {group_id}: {e}")
                members = None
                total_count = 0

            if members:
                member_ids = []
                member_data = {}
                for member in members:
                    # member dict shape may be different; use keys defensively
                    try:
                        uid = int(member.get("id") if isinstance(member.get("id"), (int, str)) else member["id"])
                    except Exception:
                        continue
                    # skip bots if flagged
                    if member.get("is_bot"):
                        continue
                    member_ids.append(uid)
                    member_data[uid] = {
                        "first_name": member.get("first_name", "") or "",
                        "last_name": member.get("last_name", "") or "",
                        "username": member.get("username", "") or "",
                        "is_bot": member.get("is_bot", False),
                        "is_admin": member.get("is_admin", False),
                    }

                gm = load_group_members()
                import time
                gm[str(group_id)] = {
                    "members": sorted(member_ids),
                    "member_data": {str(k): v for k, v in member_data.items()},
                    "last_sync": int(time.time()),
                    "sync_method": "bot_telethon",
                    "total_count": total_count or len(member_ids),
                }
                write_json(Path(GROUP_MEMBERS_FILE), gm)

                # update user_groups.json
                user_groups_data = read_json(Path(USER_GROUPS_FILE), {})
                for user_id in member_ids:
                    uid_s = str(user_id)
                    if uid_s not in user_groups_data:
                        user_groups_data[uid_s] = []
                    if group_id not in user_groups_data[uid_s]:
                        user_groups_data[uid_s].append(group_id)
                write_json(Path(USER_GROUPS_FILE), user_groups_data)

                log.info(f"Bot Telethon synced {len(member_ids)} members for group {group_id}")
                return len(member_ids)

        # 3) Fallback to aiogram (admins only)
        log.warning("Telethon clients not available or returned no members, falling back to aiogram fallback")
        return await _sync_group_members_fallback(group_id)

    except Exception as e:
        log.error(f"Failed to sync members for group {group_id}: {e}")
        return await _sync_group_members_fallback(group_id)

# ADD this new function to utils.py:
async def sync_all_groups_comprehensive() -> Dict[int, int]:
    """
    Comprehensive sync of all groups including members and admins.
    This is the main sync function that should be used for full synchronization.
    """
    results = {}
    admin_results = {}
    
    groups = load_group_ids()
    log.info(f"Starting comprehensive sync for {len(groups)} groups")
    
    for group_id in groups:
        try:
            # Sync members first
            member_count = await sync_group_members(group_id)
            results[group_id] = member_count
            
            # Small delay to avoid rate limits
            await asyncio.sleep(0.5)
            
            # Sync admins
            admin_dict = await sync_group_admins(group_id)
            admin_results[group_id] = len(admin_dict)
            
            # Another small delay
            await asyncio.sleep(0.5)
            
            log.info(f"Group {group_id}: {member_count} members, {len(admin_dict)} admins")
            
        except Exception as e:
            log.error(f"Failed to sync group {group_id}: {e}")
            results[group_id] = 0
            admin_results[group_id] = 0
    
    total_members = sum(results.values())
    total_admins = sum(admin_results.values())
    
    log.info(f"Comprehensive sync complete: {total_members} total members, {total_admins} total admins across {len(groups)} groups")
    return results


async def _sync_group_members_fallback(group_id: int) -> int:
    try:
        from config import bot
        
        admins = await bot.get_chat_administrators(group_id)
        member_ids = set()
        
        for admin in admins:
            if not admin.user.is_bot:
                member_ids.add(admin.user.id)
        
        gm = load_group_members()
        import time
        gm[str(group_id)] = {
            "members": sorted(member_ids),
            "last_sync": int(time.time()),
            "sync_method": "aiogram_fallback",
            "total_count": len(member_ids)
        }
        write_json(Path(GROUP_MEMBERS_FILE), gm)
        
        log.info(f"Fallback sync: {len(member_ids)} admin members for group {group_id}")
        return len(member_ids)
        
    except Exception as e:
        log.error(f"Fallback sync failed for group {group_id}: {e}")
        return 0

async def sync_all_groups() -> Dict[int, int]:
    """
    Sync both members and admins for all groups.
    This is a wrapper for sync_all_groups_comprehensive for backward compatibility.
    """
    return await sync_all_groups_comprehensive()

def add_user_to_group(user_id: int, group_id: int):
    update_group_member(group_id, user_id, True)
    user_groups = get_user_groups(user_id)
    if group_id not in user_groups:
        user_groups.append(group_id)
        set_user_groups(user_id, user_groups)

def get_group_member_data(group_id: int) -> Dict:
    gm = load_group_members()
    rec = gm.get(str(group_id)) or {}
    return rec.get("member_data", {})



async def notify_group_students(group_id: int, test_name: str, test_id: str) -> Tuple[List[int], List[int]]:
    """Enhanced notification with better error handling and reporting"""
    try:
        from config import bot
        from telethon_service import get_telethon_service
        
        member_ids = get_group_member_ids(group_id)
        if not member_ids:
            log.warning(f"No members found for group {group_id}")
            return [], []
        
        notified = []
        failed = []
        
        bot_username = (await bot.get_me()).username
        message = (
            f"🧪 <b>Yangi test!</b>\n\n"
            f"Test nomi: <b>{test_name}</b>\n"
            f"Test kodi: <code>{test_id}</code>\n\n"
            f"Testni boshlash uchun:\n"
            f"1. @{bot_username} botiga /start yuboring\n"
            f"2. Mavjud testlar ro'yxatidan tanlang\n\n"
            f"⏰ Test hozir faol!"
        )
        
        # Try to notify via bot first (faster)
        for user_id in member_ids:
            try:
                await bot.send_message(user_id, message)
                notified.append(user_id)
                log.debug(f"Notified user {user_id} about test {test_name}")
                
                # Small delay to avoid rate limits
                await asyncio.sleep(0.05)
                
            except Exception as e:
                failed.append(user_id)
                log.debug(f"Failed to notify user {user_id} via bot: {e}")
        
        # If many failed via bot, try Telethon as fallback
        if len(failed) > len(notified) * 0.3:  # If more than 30% failed
            log.info(f"Many bot notifications failed ({len(failed)}), trying Telethon fallback")
            
            telethon = await get_telethon_service()
            if telethon:
                telethon_message = (
                    f"🧪 Yangi test: {test_name}\n\n"
                    f"Testni boshlash uchun @{bot_username} botiga /start yuboring\n"
                    f"Test kodi: test_{test_id}"
                )
                
                telethon_notified, telethon_failed = await telethon.bulk_message_users(
                    failed, telethon_message, delay=0.5
                )
                
                notified.extend(telethon_notified)
                failed = telethon_failed
        
        success_rate = len(notified) / len(member_ids) * 100 if member_ids else 0
        log.info(f"Group {group_id} notification complete: {len(notified)}/{len(member_ids)} notified ({success_rate:.1f}% success rate)")
        
        return notified, failed
        
    except Exception as e:
        log.error(f"Failed to notify group {group_id}: {e}")
        return [], member_ids
async def _resolve_group_titles(ids):
    """Return list[(gid, title_or_id)] - resolves group IDs to titles"""
    out = []
    group_titles = load_group_titles()
    
    for gid in ids:
        try:
            gid_int = int(gid)
            title = group_titles.get(gid_int)
            
            if not title:
                # Try to get title from Telegram API
                try:
                    from config import bot
                    chat = await bot.get_chat(gid_int)
                    title = chat.title or str(gid_int)
                    
                    # Update our records
                    add_or_update_group(gid_int, title)
                    
                except Exception as e:
                    log.warning(f"Could not get title for group {gid_int}: {e}")
                    title = str(gid_int)
            
            out.append((gid_int, title))
        except (ValueError, TypeError):
            continue
    
    return out

async def detect_and_remove_kicked_users(group_id: int) -> List[int]:
    """
    Detect users who were removed from a group and clean up their data.
    Uses owner's Telethon account to get accurate member list.
    Returns list of removed user IDs.
    """
    try:
        from telethon_service import get_user_telethon_service
        
        # Get current members using owner account
        user_telethon = await get_user_telethon_service()
        if not user_telethon:
            log.warning("User Telethon not available for kicked user detection")
            return []
        
        # Get ALL current members from Telegram
        current_members, _ = await user_telethon.fetch_all_group_members(group_id)
        current_member_ids = set(member['id'] for member in current_members)
        
        # Get our stored members
        gm = load_group_members()
        stored_data = gm.get(str(group_id), {})
        stored_member_ids = set(int(uid) for uid in stored_data.get("members", []))
        
        # Find removed users (were in our data but not in current members)
        removed_users = stored_member_ids - current_member_ids
        
        if removed_users:
            log.info(f"Detected {len(removed_users)} removed users from group {group_id}: {removed_users}")
            
            # Clean up each removed user
            for user_id in removed_users:
                await cleanup_removed_user(user_id, group_id)
        
        return list(removed_users)
        
    except Exception as e:
        log.error(f"Failed to detect removed users for group {group_id}: {e}")
        return []

async def cleanup_removed_user(user_id: int, group_id: int):
    """
    Completely remove a user from all group-related data when they're kicked/removed.
    This ensures they can't access tests anymore.
    """
    try:
        log.info(f"Cleaning up removed user {user_id} from group {group_id}")
        
        # 1. Remove from group_members.json
        gm = load_group_members()
        if str(group_id) in gm:
            members = gm[str(group_id)].get("members", [])
            if user_id in members:
                members.remove(user_id)
                gm[str(group_id)]["members"] = members
            
            # Remove from member_data
            member_data = gm[str(group_id)].get("member_data", {})
            if str(user_id) in member_data:
                del member_data[str(user_id)]
                gm[str(group_id)]["member_data"] = member_data
            
            write_json(Path(GROUP_MEMBERS_FILE), gm)
            log.info(f"Removed user {user_id} from group_members.json for group {group_id}")
        
        # 2. Remove from user_groups.json
        user_groups_data = read_json(Path(USER_GROUPS_FILE), {})
        user_id_str = str(user_id)
        
        if user_id_str in user_groups_data:
            groups = user_groups_data[user_id_str]
            if group_id in groups:
                groups.remove(group_id)
                
                if groups:
                    user_groups_data[user_id_str] = groups
                else:
                    # User has no groups left, remove completely
                    del user_groups_data[user_id_str]
                
                write_json(Path(USER_GROUPS_FILE), user_groups_data)
                log.info(f"Removed group {group_id} from user {user_id}'s groups")
        
        # 3. Remove admin privileges if they had any
        admins_data = load_admins()
        group_admins = admins_data.get('group_admins', {})
        
        if user_id_str in group_admins:
            admin_groups = group_admins[user_id_str].get('groups', [])
            if group_id in admin_groups:
                admin_groups.remove(group_id)
                
                if admin_groups:
                    group_admins[user_id_str]['groups'] = admin_groups
                else:
                    # No admin groups left, remove from admins
                    del group_admins[user_id_str]
                
                admins_data['group_admins'] = group_admins
                save_admins(admins_data)
                log.info(f"Removed admin privileges for user {user_id} in group {group_id}")
        
        # 4. Log the removal for audit
        from audit import log_action
        log_action(0, "user_removed_from_group", ok=True, 
                  extra={"user_id": user_id, "group_id": group_id})
        
    except Exception as e:
        log.error(f"Failed to cleanup removed user {user_id} from group {group_id}: {e}")

async def validate_user_still_in_groups(user_id: int) -> Tuple[bool, List[int]]:
    """
    Validate if a user is still in their registered groups.
    Returns (has_valid_groups, valid_group_ids)
    """
    try:
        from config import bot
        
        # Get user's supposed groups
        user_groups = get_user_groups(user_id)
        gm = load_group_members()
        
        # Collect all groups user is supposedly in
        all_user_groups = set(user_groups)
        for gid_str, rec in gm.items():
            if user_id in rec.get("members", []):
                all_user_groups.add(int(gid_str))
        
        if not all_user_groups:
            return False, []
        
        valid_groups = []
        
        # Check each group with Telegram API
        for group_id in all_user_groups:
            try:
                member = await bot.get_chat_member(group_id, user_id)
                if member.status not in ["left", "kicked", "restricted"]:
                    valid_groups.append(group_id)
                else:
                    log.info(f"User {user_id} is no longer in group {group_id} (status: {member.status})")
                    # Clean up this invalid membership
                    await cleanup_removed_user(user_id, group_id)
            except Exception as e:
                # If we get an error (like user not found), they're not in the group
                log.info(f"User {user_id} not found in group {group_id}: {e}")
                await cleanup_removed_user(user_id, group_id)
        
        return len(valid_groups) > 0, valid_groups
        
    except Exception as e:
        log.error(f"Failed to validate user {user_id} groups: {e}")
        return False, []

async def enhanced_sync_group_members(group_id: int) -> int:
    """
    Enhanced sync that also detects and removes kicked users.
    This should replace the regular sync_group_members function when needed.
    """
    try:
        # First detect and remove kicked users
        removed_users = await detect_and_remove_kicked_users(group_id)
        if removed_users:
            log.info(f"Cleaned up {len(removed_users)} removed users from group {group_id}")
        
        # Then do normal sync
        count = await sync_group_members(group_id)
        
        return count
        
    except Exception as e:
        log.error(f"Enhanced sync failed for group {group_id}: {e}")
        return 0

async def smart_user_lookup_with_validation(user_id: int) -> bool:
    """
    Enhanced user lookup that validates membership before granting access.
    """
    try:
        # First validate if user is still in their groups
        has_valid_groups, valid_groups = await validate_user_still_in_groups(user_id)
        
        if not has_valid_groups:
            log.info(f"User {user_id} has no valid group memberships")
            return False
        
        # If user has valid groups, return True
        if valid_groups:
            log.info(f"User {user_id} validated in groups: {valid_groups}")
            return True
        
        # If no valid groups found, trigger full sync
        log.info(f"No valid groups for user {user_id}, triggering comprehensive sync")
        
        # Sync all groups with removal detection
        for group_id in load_group_ids():
            await enhanced_sync_group_members(group_id)
        
        # Check again after sync
        has_valid_groups, valid_groups = await validate_user_still_in_groups(user_id)
        
        return has_valid_groups
        
    except Exception as e:
        log.error(f"Smart user lookup with validation failed for {user_id}: {e}")
        return False

# Update the smart_user_lookup function in utils.py
async def smart_user_lookup_with_validation(user_id: int) -> bool:
    """
    Enhanced user lookup that validates membership before granting access.
    """
    try:
        # First validate if user is still in their groups
        has_valid_groups, valid_groups = await validate_user_still_in_groups(user_id)
        
        if not has_valid_groups:
            log.info(f"User {user_id} has no valid group memberships")
            return False
        
        # If user has valid groups, return True
        if valid_groups:
            log.info(f"User {user_id} validated in groups: {valid_groups}")
            return True
        
        # If no valid groups found, trigger full sync
        log.info(f"No valid groups for user {user_id}, triggering comprehensive sync")
        
        # Sync all groups with removal detection
        for group_id in load_group_ids():
            await enhanced_sync_group_members(group_id)
        
        # Check again after sync
        has_valid_groups, valid_groups = await validate_user_still_in_groups(user_id)
        
        return has_valid_groups
        
    except Exception as e:
        log.error(f"Smart user lookup with validation failed for {user_id}: {e}")
        return False

def remove_user_from_group_completely(user_id: int, group_id: int):
    """
    Completely remove user from all group-related data structures.
    This should be called when a user leaves or is kicked from a group.
    """
    try:
        # 1. Remove from group_members.json
        update_group_member(group_id, user_id, False)
        
        # 2. Remove from user_groups.json
        user_groups_data = read_json(Path(USER_GROUPS_FILE), {})
        user_id_str = str(user_id)
        
        if user_id_str in user_groups_data:
            if group_id in user_groups_data[user_id_str]:
                user_groups_data[user_id_str].remove(group_id)
                
            # If user has no groups left, remove them entirely
            if not user_groups_data[user_id_str]:
                del user_groups_data[user_id_str]
                
            write_json(Path(USER_GROUPS_FILE), user_groups_data)
        
        # 3. Remove admin privileges for this group if they had any
        remove_user_admin_privileges(user_id, group_id)
        
        log.info(f"Completely removed user {user_id} from group {group_id}")
        
    except Exception as e:
        log.error(f"Error removing user {user_id} from group {group_id}: {e}")

def remove_user_admin_privileges(user_id: int, group_id: int):
    """
    Remove admin privileges for a specific group.
    If user has no groups left, remove them from admins entirely.
    """
    try:
        admins_data = load_admins()
        group_admins = admins_data.get('group_admins', {})
        user_id_str = str(user_id)
        
        if user_id_str in group_admins:
            admin_groups = group_admins[user_id_str].get('groups', [])
            
            if group_id in admin_groups:
                admin_groups.remove(group_id)
                log.info(f"Removed admin privileges for user {user_id} in group {group_id}")
                
                # If user has no admin groups left, remove them completely
                if not admin_groups:
                    del group_admins[user_id_str]
                    log.info(f"User {user_id} has no admin groups left, removed from admins")
                else:
                    group_admins[user_id_str]['groups'] = admin_groups
                
                admins_data['group_admins'] = group_admins
                save_admins(admins_data)
                
    except Exception as e:
        log.error(f"Error removing admin privileges for user {user_id}: {e}")

async def validate_user_group_membership(user_id: int) -> List[int]:
    """
    Validate user's group membership by checking with Telegram API.
    Returns list of groups where user is actually still a member.
    """
    try:
        from config import bot
        user_groups = get_user_groups(user_id)
        valid_groups = []
        
        for group_id in user_groups:
            try:
                # Check if user is still in the group
                member = await bot.get_chat_member(group_id, user_id)
                if member.status not in ["left", "kicked"]:
                    valid_groups.append(group_id)
                else:
                    log.info(f"User {user_id} no longer in group {group_id}, status: {member.status}")
                    # Remove from all data
                    remove_user_from_group_completely(user_id, group_id)
                    
            except Exception as e:
                log.warning(f"Could not check membership for user {user_id} in group {group_id}: {e}")
                # If we can't check, assume they're still there to avoid false removals
                valid_groups.append(group_id)
        
        return valid_groups
        
    except Exception as e:
        log.error(f"Error validating group membership for user {user_id}: {e}")
        return get_user_groups(user_id)  # Return original list if validation fails

async def cleanup_invalid_memberships():
    """
    Periodic cleanup to remove users who are no longer in groups.
    This should be run periodically to catch any missed removals.
    """
    try:
        log.info("Starting cleanup of invalid memberships...")
        
        # Get all users from user_groups.json
        user_groups_data = read_json(Path(USER_GROUPS_FILE), {})
        cleanup_count = 0
        
        for user_id_str, groups in list(user_groups_data.items()):
            try:
                user_id = int(user_id_str)
                valid_groups = await validate_user_group_membership(user_id)
                
                # If valid groups differ from stored groups, update
                if set(valid_groups) != set(groups):
                    log.info(f"Updating groups for user {user_id}: {groups} -> {valid_groups}")
                    
                    if valid_groups:
                        user_groups_data[user_id_str] = valid_groups
                    else:
                        del user_groups_data[user_id_str]
                    
                    cleanup_count += 1
                    
            except Exception as e:
                log.error(f"Error cleaning up user {user_id_str}: {e}")
        
        # Save updated data
        if cleanup_count > 0:
            write_json(Path(USER_GROUPS_FILE), user_groups_data)
            log.info(f"Cleanup complete: updated {cleanup_count} users")
        else:
            log.info("Cleanup complete: no changes needed")
            
    except Exception as e:
        log.error(f"Error during membership cleanup: {e}")


async def check_user_current_admin_status(user_id: int) -> List[int]:
    """
    Check user's CURRENT admin status by validating with Telegram API.
    Returns list of groups where user is currently an admin.
    """
    try:
        from config import bot
        
        # Get stored admin groups
        stored_admin_groups = get_user_admin_groups(user_id)
        current_admin_groups = []
        
        for group_id in stored_admin_groups:
            try:
                # Check current status in Telegram
                member = await bot.get_chat_member(group_id, user_id)
                if member.status in ["administrator", "creator"]:
                    current_admin_groups.append(group_id)
                else:
                    log.info(f"User {user_id} no longer admin in group {group_id}, current status: {member.status}")
                    # Remove admin privileges
                    remove_user_admin_privileges(user_id, group_id)
                    
            except Exception as e:
                log.warning(f"Could not check admin status for user {user_id} in group {group_id}: {e}")
                # If we can't check, don't assume they're still admin
        
        return current_admin_groups
        
    except Exception as e:
        log.error(f"Error checking admin status for user {user_id}: {e}")
        return []

async def get_unreachable_users_info(group_id: int, failed_user_ids: List[int]) -> List[Dict]:
    try:
        member_data = get_group_member_data(group_id)
        unreachable = []
        
        for user_id in failed_user_ids:
            user_info = member_data.get(str(user_id), {})
            unreachable.append({
                'id': user_id,
                'first_name': user_info.get('first_name', 'Unknown'),
                'last_name': user_info.get('last_name', ''),
                'username': user_info.get('username', ''),
                'full_name': f"{user_info.get('first_name', 'Unknown')} {user_info.get('last_name', '')}".strip()
            })
        
        return unreachable
        
    except Exception as e:
        log.error(f"Failed to get unreachable user info: {e}")
        return []

def load_students() -> Dict[str, dict]:
    return read_json(Path(STUDENTS_FILE), {})

def save_students(data: Dict[str, dict]):
    write_json(Path(STUDENTS_FILE), data)

def save_student_data(user_id: int, payload: dict):
    students = load_students()
    base = students.get(str(user_id), {})
    base.update(payload or {})
    students[str(user_id)] = base
    save_students(students)

def get_user_groups(user_id: int) -> List[int]:
    m = read_json(Path(USER_GROUPS_FILE), {})
    ids = m.get(str(user_id), [])
    try:
        return [int(x) for x in ids]
    except Exception:
        return []

def set_user_groups(user_id: int, groups: List[int]):
    m = read_json(Path(USER_GROUPS_FILE), {})
    m[str(user_id)] = sorted(set(int(x) for x in groups))
    write_json(Path(USER_GROUPS_FILE), m)

# Add this fix to your utils.py file in the get_all_students_with_groups function

def get_all_students_with_groups() -> List[Dict]:
    students = []
    gm = load_group_members()
    group_titles = load_group_titles()
    
    all_users = {}
    
    for group_id_str, group_data in gm.items():
        group_id = int(group_id_str)
        member_data = group_data.get("member_data", {})
        
        for user_id_str, user_info in member_data.items():
            user_id = int(user_id_str)
            if user_id not in all_users:
                all_users[user_id] = {
                    'groups': [],
                    'info': user_info
                }
            all_users[user_id]['groups'].append(group_id)
    
    user_groups_data = read_json(Path(USER_GROUPS_FILE), {})
    for user_id_str, group_ids in user_groups_data.items():
        user_id = int(user_id_str)
        if user_id not in all_users:
            all_users[user_id] = {
                'groups': [int(gid) for gid in group_ids],
                'info': {}
            }
    
    student_data = load_students()
    
    for user_id, data in all_users.items():
        user_info = data['info']
        user_groups = data['groups']
        
        group_names = []
        for gid in user_groups:
            title = group_titles.get(gid, f"Group {gid}")
            group_names.append(title)
        
        test_info = student_data.get(str(user_id), {})
        
        name_parts = []
        if user_info.get('first_name'):
            name_parts.append(user_info['first_name'])
        if user_info.get('last_name'):
            name_parts.append(user_info['last_name'])
        
        display_name = " ".join(name_parts) if name_parts else test_info.get("full_name", "Unknown")
        
        students.append({
            "user_id": user_id,
            "name": display_name,
            "username": user_info.get('username', ''),
            "groups": group_names,
            "last_test": test_info.get("last_test_id"),
            "last_score": test_info.get("last_score"),
            "last_activity": test_info.get("finished_at"),
        })
    
    # FIX: Handle None values in sorting by using 'or 0' to provide default
    return sorted(students, key=lambda x: x.get("last_activity") or 0, reverse=True)

def load_admins() -> Dict[str, dict]:
    p = Path(DATA_DIR) / "admins.json"
    d = read_json(p, {"owner_id": OWNER_ID, "admins": {}})
    return d

def save_admins(d: Dict[str, dict]):
    p = Path(DATA_DIR) / "admins.json"
    write_json(p, d)

# Add to utils.py
def get_student_admins(user_id: int) -> List[int]:
    """Get admin IDs who manage groups where this student is a member"""
    try:
        # Get student's groups
        user_groups = get_user_groups(user_id)
        gm = load_group_members()
        
        # Also check group_members.json for groups
        for gid_str, rec in gm.items():
            try:
                gid = int(gid_str)
                if user_id in rec.get("members", []):
                    if gid not in user_groups:
                        user_groups.append(gid)
            except (ValueError, TypeError):
                continue
        
        if not user_groups:
            return []
        
        # Get all admins for these groups
        admins_data = load_admins()
        group_admins = admins_data.get('group_admins', {})
        
        relevant_admins = []
        for admin_id_str, admin_data in group_admins.items():
            admin_groups = admin_data.get('groups', [])
            # Check if this admin manages any of the student's groups
            if any(g in user_groups for g in admin_groups):
                relevant_admins.append(int(admin_id_str))
        
        return relevant_admins
        
    except Exception as e:
        log.error(f"Error getting student admins: {e}")
        return []

def test_path(test_id: str) -> Path:
    p = Path(TESTS_DIR) / f"test_{test_id}.json"
    ensure_dir(p.parent)
    return p

def get_all_tests() -> List[Path]:
    return list(Path(TESTS_DIR).glob("test_*.json"))

def read_test(test_id: str) -> dict:
    p = test_path(test_id)
    return read_json(p, {})

def write_test(test_id: str, obj: dict):
    p = test_path(test_id)
    write_json(p, obj)

def add_test_index(test_id: str, name: str):
    obj = read_test(test_id) or {}
    obj.setdefault("test_id", test_id)
    if name:
        obj["test_name"] = name
    obj.setdefault("questions", [])
    obj.setdefault("answers", {})
    obj.setdefault("references", {})
    obj.setdefault("groups", [])
    if "group_id" not in obj:
        obj["group_id"] = None
    write_test(test_id, obj)

def save_test_content(test_id: str, content: dict):
    obj = read_test(test_id) or {"test_id": test_id}
    for k in ("test_name", "questions", "answers", "references"):
        if k in content:
            obj[k] = content[k]
    obj.setdefault("groups", [])
    obj.setdefault("group_id", None)
    write_test(test_id, obj)

def load_tests_index() -> dict:
    active = set(get_active_tests())
    out = {"tests": {}}
    for p in get_all_tests():
        try:
            data = json.loads(p.read_text(encoding="utf-8"))
        except Exception:
            continue
        tid = data.get("test_id") or p.stem.replace("test_", "")
        name = data.get("test_name") or "Test"
        groups = data.get("groups") or []
        gid = data.get("group_id")
        if gid and not groups:
            groups = [str(gid)]
        out["tests"][tid] = {
            "name": name,
            "active": tid in active,
            "groups": [str(g) for g in groups],
        }
    return out

def save_tests_index(idx: dict):
    return

def get_active_tests() -> List[str]:
    data = read_json(Path(ACTIVE_TEST_FILE), {"active_tests": []})
    return list(dict.fromkeys(data.get("active_tests", [])))

def set_active_test(test_id: str):
    data = read_json(Path(ACTIVE_TEST_FILE), {"active_tests": []})
    arr = list(data.get("active_tests", []))
    if test_id not in arr:
        arr.append(test_id)
    data["active_tests"] = arr
    write_json(Path(ACTIVE_TEST_FILE), data)

def remove_active_test(test_id: str):
    data = read_json(Path(ACTIVE_TEST_FILE), {"active_tests": []})
    arr = [x for x in data.get("active_tests", []) if x != test_id]
    data["active_tests"] = arr
    write_json(Path(ACTIVE_TEST_FILE), data)

def set_test_active(test_id: str, active: bool):
    if active:
        set_active_test(test_id)
    else:
        remove_active_test(test_id)

def assign_test_groups(test_id: str, groups: List[int]):
    obj = read_test(test_id) or {"test_id": test_id}
    norm = [int(x) for x in groups if str(x).strip()]
    obj["groups"] = sorted(set(norm))
    obj["group_id"] = (norm[0] if norm else None)
    write_test(test_id, obj)

def tests_for_group(group_id: int, only_active: bool = True) -> List[dict]:
    """
    Guruh bo'yicha testlar ro'yxati.
    only_active=True bo'lsa:
      - test global aktiv bo'lishi shart
      - agar test.active_groups mavjud bo'lsa, group_id o'shanda bo'lishi shart
    """
    act = set(get_active_tests()) if only_active else None
    out: List[dict] = []

    for p in get_all_tests():
        try:
            data = json.loads(p.read_text(encoding="utf-8"))
        except Exception:
            continue

        tid = data.get("test_id") or p.stem.replace("test_", "")

        # Test tayinlangan guruhlar
        grps: List[int] = []
        for g in (data.get("groups") or []):
            try:
                grps.append(int(g))
            except Exception:
                pass
        gid = data.get("group_id")
        if gid and not grps:
            try:
                grps = [int(gid)]
            except Exception:
                grps = []

        if int(group_id) not in grps:
            continue

        if only_active:
            # Global aktiv shart
            if act is not None and tid not in act:
                continue
            # Guruh bo'yicha aktiv filtri
            active_groups = set()
            for g in (data.get("active_groups") or []):
                try:
                    active_groups.add(int(g))
                except Exception:
                    pass
            if active_groups and (int(group_id) not in active_groups):
                continue

        data["test_id"] = tid
        out.append(data)

    return out



def available_tests_for_user(user_id: int) -> List[dict]:
    """
    Foydalanuvchiga ko'rinadigan (faol) testlar ro'yxati.
    -- Global aktiv: ACTIVE_TEST_FILE dagi testlar
    -- Guruh bo'yicha aktiv: test.json ichidagi "active_groups"
       * Agar "active_groups" bo'sh/yo'q bo'lsa, tayinlangan barcha guruhlar uchun aktiv hisoblanadi.
       * Agar "active_groups" bor bo'lsa, faqat o'sha guruhlar uchun aktiv.
    """
    user_groups = set(get_user_groups(user_id))

    # group_members.json bo'yicha ham tekshiramiz
    gm = load_group_members()
    for gid_str, rec in gm.items():
        try:
            gid = int(gid_str)
            members = set(int(x) for x in rec.get("members", []) or [])
            if user_id in members:
                user_groups.add(gid)
        except (ValueError, TypeError):
            continue

    if not user_groups:
        log.warning(f"User {user_id} has no group memberships")
        return []

    log.info(f"User {user_id} is member of groups: {sorted(user_groups)}")

    out: List[dict] = []
    active_tests = set(get_active_tests())

    for p in get_all_tests():
        try:
            data = json.loads(p.read_text(encoding="utf-8"))
        except Exception:
            continue

        tid = data.get("test_id") or p.stem.replace("test_", "")

        # 1) Global aktiv filtri
        if tid not in active_tests:
            continue

        # 2) Test tayinlangan guruhlar
        test_groups = set()
        for g in (data.get("groups") or []):
            try:
                test_groups.add(int(g))
            except (ValueError, TypeError):
                pass
        if data.get("group_id"):
            try:
                test_groups.add(int(data["group_id"]))
            except (ValueError, TypeError):
                pass
        if not test_groups:
            continue

        # 3) Guruh bo'yicha faollik
        active_groups = set()
        for g in (data.get("active_groups") or []):
            try:
                active_groups.add(int(g))
            except Exception:
                pass

        effective_groups = (test_groups & active_groups) if active_groups else test_groups

        if user_groups & effective_groups:
            data["test_id"] = tid
            out.append(data)
            log.info(f"Test {tid} available for user {user_id} (effective groups: {sorted(effective_groups)})")

    return out


_Q_HEADER = re.compile(r"^\s*(\d+)[\.\)]\s*(.+)$")
_OPT_LINE = re.compile(r"^\s*([A-Da-d])[\.\)]\s*(.+)$")
_ANS_PAIR = re.compile(r"^\s*(\d+)\s*[\.\)\-]?\s*([A-Da-d])\s*$")
_ANS_TOK = re.compile(r"^\s*(\d+)([A-Da-d])\s*$")
_REF_LINE = re.compile(r"^\s*(\d+)[\.\)]\s*(.+)$")






def parse_answers_from_text(block: str) -> Dict[str, str]:
    """Enhanced answer parser that handles various formats"""
    ans: Dict[str, str] = {}
    
    # Handle different answer formats
    for line in block.splitlines():
        line = line.strip()
        if not line:
            continue
            
        # Try different patterns
        # Pattern 1: "1.a" or "1)a" 
        m = re.match(r'^\s*(\d+)\s*[\.\\)]\s*([A-Da-d])\s*$', line)
        if m:
            ans[str(int(m.group(1)))] = m.group(2).upper()
            continue
            
        # Pattern 2: "1-a" or "1:a"
        m = re.match(r'^\s*(\d+)\s*[-:]\s*([A-Da-d])\s*$', line)
        if m:
            ans[str(int(m.group(1)))] = m.group(2).upper()
            continue
            
        # Pattern 3: "1a" (no separator)
        m = re.match(r'^\s*(\d+)([A-Da-d])\s*$', line)
        if m:
            ans[str(int(m.group(1)))] = m.group(2).upper()
            continue
            
        # Pattern 4: Split by commas or spaces for multiple answers
        tokens = re.split(r'[,\s]+', line)
        for token in tokens:
            token = token.strip()
            if not token:
                continue
                
            # Try each token with the patterns above
            for pattern in [
                r'^(\d+)[\.\\)]?([A-Da-d])$',
                r'^(\d+)[-:]([A-Da-d])$'
            ]:
                m = re.match(pattern, token)
                if m:
                    ans[str(int(m.group(1)))] = m.group(2).upper()
                    break
    
    return ans


def parse_questions_from_text(text: str) -> List[dict]:
    """
    Robust parser for numbered questions with options A–D that:
      - Preserves ALL lines between the question header and the first option (keeps code fences, blank lines).
      - Treats a line like '2)' or '3.' as a new question ONLY if options (A–D) appear shortly after (lookahead).
      - Allows multi-line options until the next option or a new question header (with lookahead guard).

    Returns a list of dicts: {"index": int, "text": str, "options": {"A": ..., "B": ...}}
    """
    text = (text or "").replace("\r\n", "\n").replace("\r", "\n")
    lines = [ln.rstrip() for ln in text.split("\n")]

    q_re    = re.compile(r'^\s*(\d+)\s*[.)]\s*(.*)$')            # "3) ..." or "3. ..."
    opt_re  = re.compile(r'^\s*([A-Da-d])\s*[.)]\s+(.*)$')        # "A) ..." (note the required space)
    fence_re = re.compile(r'^\s*```')                             # code-fence toggle

    items: List[dict] = []
    cur: Optional[dict] = None
    in_code = False

    def has_options_ahead(start_idx: int, window: int = 12) -> bool:
        end = min(len(lines), start_idx + window)
        for j in range(start_idx, end):
            if opt_re.match(lines[j]):
                return True
        return False

    def push_current():
        nonlocal cur
        if not cur:
            return
        # valid only if it has at least 2 options and some text
        if (cur.get("text") or "").strip() and len(cur.get("options", {})) >= 2:
            items.append(cur)
        cur = None

    i = 0
    # main loop
    while i < len(lines):
        ln = lines[i]

        # toggle code fence
        if fence_re.match(ln):
            if cur is None:
                # ignore stray fences before first question
                i += 1
                continue
            if cur and not cur["options"]:
                # still in question text
                cur["text"] = (cur["text"] + ("\n" if cur["text"] else "") + ln)
            else:
                # inside options: append to last option if any
                if cur["options"]:
                    last_key = sorted(cur["options"].keys())[-1]
                    cur["options"][last_key] += ("\n" + ln)
            in_code = not in_code
            i += 1
            continue

        # New question header?
        m_q = q_re.match(ln) if not in_code else None
        if m_q:
            # Only treat as a new question if there are options in the near future
            if has_options_ahead(i + 1):
                push_current()
                idx = int(m_q.group(1))
                first = (m_q.group(2) or "").strip()
                cur = {"index": idx, "text": first, "options": {}}

                # Collect subsequent lines into question text until options start
                i += 1
                while i < len(lines):
                    probe = lines[i]
                    if opt_re.match(probe):
                        break
                    # If another numeric header appears, it’s only a real split if options follow it too
                    if q_re.match(probe) and has_options_ahead(i + 1):
                        break
                    # keep structure (include blanks)
                    if cur["text"]:
                        cur["text"] += ("\n" + probe)
                    else:
                        cur["text"] = probe
                    i += 1
                continue
            else:
                # looks like numbered text inside question/code — treat as continuation
                if cur:
                    cur["text"] = (cur["text"] + ("\n" if cur["text"] else "") + ln)
                i += 1
                continue

        # Option header?
        m_opt = opt_re.match(ln) if not in_code else None
        if m_opt and cur:
            key = m_opt.group(1).upper()
            val = (m_opt.group(2) or "").strip()

            # Collect multi-line option body
            i += 1
            parts = [val] if val else []
            while i < len(lines):
                nxt = lines[i]
                if opt_re.match(nxt):
                    break
                if q_re.match(nxt) and has_options_ahead(i + 1):
                    break
                parts.append(nxt)  # keep raw; trim later
                i += 1
            cur["options"][key] = "\n".join(s.rstrip() for s in parts).strip()
            continue

        # Otherwise, continuation of question text (before options)
        if cur:
            cur["text"] = (cur["text"] + ("\n" if cur["text"] else "") + ln)

        i += 1

    # flush last
    push_current()

    # Re-number sequentially for consistency
    for j, q in enumerate(items, 1):
        q["index"] = j

    return items


def parse_questions_from_text_v2(text: str) -> List[dict]:
    """
    Fallback parser that reuses your existing parse_questions_from_text_enhanced(...).
    Kept separate so smart_parse_questions can compare results and pick the better one.
    """
    try:
        return parse_questions_from_text_enhanced(text)
    except Exception:
        return []
    
def smart_parse_questions(text: str) -> List[dict]:
    """
    Try the robust parser first, fall back to v2 (enhanced) and pick the better result.
    'Better' = more questions that each have text + at least 2 options.
    """
    def score(lst: List[dict]) -> Tuple[int, int]:
        # (#valid, total)
        valid = 0
        for q in lst or []:
            if (q.get("text") or "").strip() and len((q.get("options") or {})) >= 2:
                valid += 1
        return valid, len(lst or [])

    a = parse_questions_from_text(text)
    b = parse_questions_from_text_v2(text)

    sa = score(a)
    sb = score(b)

    if sa > sb:
        return a
    if sb > sa:
        return b
    # tie-breaker: prefer the one with more lines in first question text (tends to keep code blocks)
    def first_q_len(lst: List[dict]) -> int:
        if not lst:
            return 0
        return len((lst[0].get("text") or "").splitlines())
    return a if first_q_len(a) >= first_q_len(b) else b


def smart_parse_questions_v2(text: str) -> List[dict]:
    """Compatibility alias."""
    return smart_parse_questions(text)



def parse_questions_from_text_enhanced(block: str) -> List[dict]:
    """Fixed parser that properly handles multi-line questions with code blocks"""
    lines = [ln.rstrip() for ln in block.splitlines()]
    items = []
    i = 0
    
    while i < len(lines):
        line = lines[i].strip()
        
        # Skip empty lines
        if not line:
            i += 1
            continue
        
        # Check for question header (1. 2. etc)
        q_match = re.match(r'^(\d+)\s*[\.\\)]\s*(.*)$', line)
        
        if q_match:
            qnum = int(q_match.group(1))
            question_text_parts = []
            
            # Add initial question text if present
            initial_text = q_match.group(2).strip()
            if initial_text:
                question_text_parts.append(initial_text)
            
            i += 1
            
            # Collect ALL question text until we find options
            while i < len(lines):
                current_line = lines[i]
                
                # Check if this line starts an option (a) b) c) d))
                if re.match(r'^\s*[A-Da-d]\s*[\.\\)]\s*', current_line):
                    break
                
                # Check if this is the next question
                next_q = re.match(r'^\s*(\d+)\s*[\.\\)]\s*', current_line)
                if next_q and int(next_q.group(1)) > qnum:
                    break
                
                # Add this line to question text
                question_text_parts.append(current_line.rstrip())
                i += 1
            
            # Join question text preserving structure
            full_question = '\n'.join(question_text_parts).strip()
            
            # Collect options
            options = {}
            while i < len(lines):
                opt_line = lines[i].strip()
                opt_match = re.match(r'^([A-Da-d])\s*[\.\\)]\s*(.*)$', opt_line)
                
                if opt_match:
                    key = opt_match.group(1).upper()
                    value = opt_match.group(2).strip()
                    
                    # Collect multi-line option text
                    i += 1
                    option_parts = [value] if value else []
                    
                    while i < len(lines):
                        next_line = lines[i].strip()
                        
                        # Stop if next option or question
                        if (re.match(r'^[A-Da-d]\s*[\.\\)]\s*', next_line) or 
                            re.match(r'^\d+\s*[\.\\)]\s*', next_line)):
                            break
                        
                        if next_line:
                            option_parts.append(next_line)
                        i += 1
                    
                    options[key] = ' '.join(option_parts).strip()
                else:
                    break
            
            # Only add if we have valid question and options
            if full_question and len(options) >= 2:
                items.append({
                    "index": qnum,
                    "text": full_question,
                    "options": options
                })
        else:
            i += 1
    
    return items






# Replace the existing parse_combined_file function
def parse_combined_file_enhanced(file_path: str) -> Optional[dict]:
    """Enhanced DOCX parser with better code block handling"""
    try:
        doc = Document(file_path)
    except Exception as e:
        raise ValueError(f"Cannot open DOCX: {e}")

    qtxt, atxt, rtxt = _split_sections_from_doc(doc)
    
    # Use enhanced parser
    questions = smart_parse_questions_v2(qtxt)
    answers = parse_answers_from_text(atxt)
    refs = parse_references_from_text(rtxt)

    if not questions:
        raise ValueError("No questions parsed. Ensure 'Savollar' section and proper numbering.")
    
    # Validate answers
    for k, v in answers.items():
        if v not in {"A", "B", "C", "D"}:
            raise ValueError(f"Answer {k} has invalid option '{v}'")

    # Generate test name from first question or filename
    test_name = Path(file_path).stem
    if questions and questions[0]["text"]:
        first_q_text = questions[0]["text"][:50].replace('\n', ' ').strip()
        if first_q_text:
            test_name = first_q_text

    return {
        "test_name": test_name,
        "questions": questions,
        "answers": answers,
        "references": refs,
    }

def parse_docx_bytes_enhanced(raw: bytes) -> Tuple[str, dict]:
    """Enhanced DOCX parser wrapper"""
    with tempfile.NamedTemporaryFile(suffix=".docx", delete=False) as tmp:
        tmp.write(raw)
        tmp.flush()
        tmp_path = tmp.name
    try:
        parsed = parse_combined_file_enhanced(tmp_path)
        name = parsed.get("test_name") or "Test"
        return name, parsed
    finally:
        try:
            os.unlink(tmp_path)
        except Exception:
            pass

def parse_references_from_text(block: str) -> Dict[str, str]:
    """Enhanced reference parser that preserves code formatting"""
    refs: Dict[str, str] = {}
    current_ref = None
    current_text = []
    
    for line in block.splitlines():
        # Check if this is a new reference number
        m = re.match(r'^\s*(\d+)[\.\\)]\s*(.*)$', line)
        if m:
            # Save previous reference if exists
            if current_ref and current_text:
                refs[str(current_ref)] = ' '.join(current_text).strip()
            
            # Start new reference
            current_ref = int(m.group(1))
            remaining = m.group(2).strip()
            current_text = [remaining] if remaining else []
        elif current_ref and line.strip():
            # Continue current reference
            current_text.append(line.strip())
    
    # Don't forget the last reference
    if current_ref and current_text:
        refs[str(current_ref)] = ' '.join(current_text).strip()
    
    return refs

def validate_parsed_test(parsed_data: dict) -> Tuple[bool, str]:
    """Validate parsed test data"""
    questions = parsed_data.get("questions", [])
    answers = parsed_data.get("answers", {})
    
    if not questions:
        return False, "No questions found"
    
    # Check that all questions have at least 2 options
    for q in questions:
        opts = q.get("options", {})
        if len(opts) < 2:
            return False, f"Question {q.get('index')} has less than 2 options"
    
    # Check that answers exist for all questions
    for q in questions:
        idx = str(q.get("index"))
        if idx not in answers:
            return False, f"No answer found for question {idx}"
        if answers[idx] not in ["A", "B", "C", "D"]:
            return False, f"Invalid answer '{answers[idx]}' for question {idx}"
    
    return True, "Valid"

def _split_sections_from_doc(doc: Document) -> Tuple[str, str, str]:
    """Enhanced section splitter that preserves ALL content including code"""
    sections = {"questions": [], "answers": [], "refs": []}
    current_section = "questions"
    
    # More flexible header patterns
    headers = {
        "questions": re.compile(r'^\s*\**\s*Savollar\s*:?\s*\**\s*$', re.IGNORECASE),
        "answers": re.compile(r'^\s*\**\s*Javoblar\s*:?\s*\**\s*$', re.IGNORECASE),
        "refs": re.compile(r'^\s*\**\s*Izohlar\s*:?\s*\**\s*$', re.IGNORECASE)
    }
    
    for para in doc.paragraphs:
        text = para.text
        
        # Check if this is a section header
        header_found = False
        for section, pattern in headers.items():
            if pattern.match(text.strip()):
                current_section = section if section != "refs" else "refs"
                header_found = True
                break
        
        if header_found:
            continue
        
        # Add to current section, preserving ALL content
        if current_section == "questions":
            sections["questions"].append(text)
        elif current_section == "answers":
            sections["answers"].append(text)
        else:
            sections["refs"].append(text)
    
    # Join with newlines to preserve structure
    return (
        "\n".join(sections["questions"]),
        "\n".join(sections["answers"]),
        "\n".join(sections["refs"])
    )




# Additional helper to validate and fix incomplete questions


def parse_combined_file(file_path: str) -> Optional[dict]:
    """
    Parse a single DOCX that contains three sections:
      - Savollar: numbered questions with options A–D
      - Javoblar: mapping of question number -> correct option letter
      - Izohlar (optional): references/notes

    This version intentionally routes question parsing through the "smart" parser and
    then post-fixes questions to avoid truncation of multi-line/code questions.
    """
    try:
        doc = Document(file_path)
    except Exception as e:
        raise ValueError(f"Cannot open DOCX: {e}")

    qtxt, atxt, rtxt = _split_sections_from_doc(doc)

    # Use the safer pipeline to avoid cutting off question text (e.g., Q3).
    questions = smart_parse_questions(qtxt)
    questions = validate_and_fix_questions(questions, qtxt)

    answers = parse_answers_from_text(atxt)
    refs = parse_references_from_text(rtxt)

    if not questions:
        raise ValueError(
            "No questions parsed. Ensure 'Savollar' section exists and questions are like '1) ...' "
            "with options A–D on their own lines."
        )

    # Validate answers format
    for k, v in answers.items():
        if v not in {"A", "B", "C", "D"}:
            raise ValueError(f"Answer {k} has invalid option '{v}'")

    # Derive a test name (fallback to filename)
    test_name = Path(file_path).stem
    if questions and (questions[0].get("text") or "").strip():
        # Use a snippet from the first question as a friendlier name (optional)
        snippet = questions[0]["text"].strip().splitlines()[0]
        if snippet:
            test_name = snippet[:80]

    return {
        "test_name": test_name,
        "questions": questions,
        "answers": answers,
        "references": refs,
    }


def validate_and_fix_questions(questions: List[dict], raw_q_text: str) -> List[dict]:
    """
    Light normalizer:
      - Keep only options A..D (uppercased), drop extras.
      - Drop questions with <2 options or empty text.
      - Trim whitespace but preserve internal newlines.
      - Renumber sequentially.
    """
    fixed: List[dict] = []
    for q in questions or []:
        text = (q.get("text") or "").strip("\n")
        opts = {k.upper(): (v or "").strip("\n") for k, v in (q.get("options") or {}).items()}
        # keep only A-D in order
        clean_opts = {}
        for k in ("A", "B", "C", "D"):
            if k in opts and opts[k].strip():
                clean_opts[k] = opts[k]

        if text and len(clean_opts) >= 2:
            fixed.append({"index": q.get("index") or 0, "text": text, "options": clean_opts})

    # renumber 1..N
    for i, q in enumerate(fixed, 1):
        q["index"] = i

    return fixed


def parse_docx_bytes(raw: bytes) -> Tuple[str, dict]:
    """Compat wrapper: eski nom → parse_docx_bytes_enhanced."""
    return parse_docx_bytes_enhanced(raw)

def score_user_answers(user_answers: Dict[str, str], correct: Dict[str, str]) -> Tuple[int, int]:
    total = len(correct or {})
    if total == 0:
        return 0, 0
    ok = 0
    for k, v in user_answers.items():
        if str(k) in correct and str(v).upper() == str(correct[str(k)]).upper():
            ok += 1
    return ok, total


def validate_test_id(test_id: str) -> bool:
    """
    Test ID uchun yagona, qat’iy validator.
    Harf/raqam, chiziqcha (-) va pastki chiziq (_) ruxsat.
    Uzunligi: 1..100
    """
    if not isinstance(test_id, str) or not test_id or len(test_id) > 100:
        return False
    return bool(re.match(r'^[A-Za-z0-9_-]+$', test_id))



_session_locks = {}
_session_lock_main = asyncio.Lock()

async def get_session_lock(user_id: int, test_id: str) -> asyncio.Lock:
    """Get a lock specific to this user's test session"""
    key = f"{user_id}:{test_id}"
    async with _session_lock_main:
        if key not in _session_locks:
            _session_locks[key] = asyncio.Lock()
        return _session_locks[key]

def get_session_path(user_id: int, test_id: str) -> Path:
    """Get the path for a specific session file"""
    session_dir = Path("data/sessions")
    session_dir.mkdir(parents=True, exist_ok=True)
    return session_dir / f"session_{user_id}_{test_id}.json"

async def save_test_session_safe(user_id: int, test_id: str, session_data: dict) -> bool:
    """Thread-safe session saving with error handling"""
    try:
        session_lock = await get_session_lock(user_id, test_id)
        async with session_lock:
            # Prepare data for JSON serialization
            clean_data = session_data.copy()
            clean_data['last_updated'] = int(time.time())
            
            # Handle excluded_options conversion
            if 'excluded_options' in clean_data:
                clean_excluded = {}
                for k, v in clean_data['excluded_options'].items():
                    if isinstance(v, set):
                        clean_excluded[k] = list(v)
                    elif isinstance(v, list):
                        clean_excluded[k] = v
                    else:
                        clean_excluded[k] = []
                clean_data['excluded_options'] = clean_excluded
            
            # Write to temporary file first, then move (atomic operation)
            session_path = get_session_path(user_id, test_id)
            temp_path = session_path.with_suffix('.tmp')
            
            with open(temp_path, 'w', encoding='utf-8') as f:
                json.dump(clean_data, f, ensure_ascii=False, indent=2)
            
            # Atomic move
            temp_path.replace(session_path)
            return True
            
    except Exception as e:
        log.error(f"Failed to save session for user {user_id}, test {test_id}: {e}")
        return False

async def load_test_session_safe(user_id: int, test_id: str) -> Optional[dict]:
    """Thread-safe session loading with error handling"""
    try:
        session_lock = await get_session_lock(user_id, test_id)
        async with session_lock:
            session_path = get_session_path(user_id, test_id)
            
            if not session_path.exists():
                return None
            
            with open(session_path, 'r', encoding='utf-8') as f:
                session_data = json.load(f)
            
            # Convert excluded_options back to the expected format
            if 'excluded_options' in session_data:
                for k, v in session_data['excluded_options'].items():
                    if isinstance(v, list):
                        session_data['excluded_options'][k] = v
            
            return session_data
            
    except Exception as e:
        log.error(f"Failed to load session for user {user_id}, test {test_id}: {e}")
        return None

async def delete_test_session_safe(user_id: int, test_id: str) -> bool:
    """Thread-safe session deletion"""
    try:
        session_lock = await get_session_lock(user_id, test_id)
        async with session_lock:
            session_path = get_session_path(user_id, test_id)
            
            if session_path.exists():
                session_path.unlink()
                return True
            return False
            
    except Exception as e:
        log.error(f"Failed to delete session for user {user_id}, test {test_id}: {e}")
        return False

def get_user_sessions_safe(user_id: int) -> List[Dict]:
    """Get all sessions for a user"""
    try:
        session_dir = Path("data/sessions")
        if not session_dir.exists():
            return []
        
        sessions = []
        pattern = f"session_{user_id}_*.json"
        
        for session_file in session_dir.glob(pattern):
            try:
                with open(session_file, 'r', encoding='utf-8') as f:
                    session_data = json.load(f)
                
                test_id = session_data.get("active_test_id")
                if test_id:
                    # Get test info
                    test = read_test(test_id)
                    if test:
                        current_q = session_data.get("current_q", 1)
                        total_q = session_data.get("total_q", 0)
                        
                        sessions.append({
                            "test_id": test_id,
                            "test_name": test.get("test_name", "Test"),
                            "progress": f"{current_q}/{total_q}",
                            "started_at": session_data.get("started_at", 0),
                            "last_updated": session_data.get("last_updated", 0)
                        })
                        
            except Exception as e:
                log.warning(f"Error reading session file {session_file}: {e}")
                continue
        
        # Sort by last updated
        sessions.sort(key=lambda x: x.get("last_updated", 0), reverse=True)
        return sessions
        
    except Exception as e:
        log.error(f"Failed to get user sessions for {user_id}: {e}")
        return []

# Improved error handling for student operations
async def handle_student_error(cb_or_msg, error: Exception, context: str = "operation"):
    """Centralized error handling for student operations"""
    log.error(f"Student {context} error: {error}")
    
    user_friendly_message = "Xatolik yuz berdi. Iltimos qaytadan urinib ko'ring."
    
    # Specific error handling
    if "Test topilmadi" in str(error):
        user_friendly_message = "Test topilmadi yoki o'chirilgan. /start buyrug'ini yuboring."
    elif "permission" in str(error).lower():
        user_friendly_message = "Bu test uchun ruxsatingiz yo'q."
    elif "session" in str(error).lower():
        user_friendly_message = "Sessiya xatoligi. /start buyrug'ini yuboring."
    
    try:
        if hasattr(cb_or_msg, 'answer'):  # CallbackQuery
            await cb_or_msg.answer(user_friendly_message, show_alert=True)
        else:  # Message
            await cb_or_msg.reply(user_friendly_message)
    except Exception as e:
        log.error(f"Failed to send error message: {e}")

        # Add these wrapper functions to utils.py to maintain compatibility with student_handlers.py

def save_test_session(user_id: int, test_id: str, session_data: dict) -> bool:
    """Synchronous wrapper for save_test_session_safe"""
    import asyncio
    try:
        loop = asyncio.get_event_loop()
        if loop.is_running():
            # If called from async context, create task
            future = asyncio.ensure_future(save_test_session_safe(user_id, test_id, session_data))
            return True  # Return immediately, save happens in background
        else:
            # If called from sync context, run until complete
            return loop.run_until_complete(save_test_session_safe(user_id, test_id, session_data))
    except Exception as e:
        log.error(f"Error in save_test_session wrapper: {e}")
        # Fallback to simple file write
        try:
            session_path = get_session_path(user_id, test_id)
            session_path.parent.mkdir(parents=True, exist_ok=True)
            with open(session_path, 'w', encoding='utf-8') as f:
                json.dump(session_data, f, ensure_ascii=False, indent=2)
            return True
        except Exception as fallback_error:
            log.error(f"Fallback save also failed: {fallback_error}")
            return False

def load_test_session(user_id: int, test_id: str) -> Optional[dict]:
    """Synchronous wrapper for load_test_session_safe"""
    import asyncio
    try:
        loop = asyncio.get_event_loop()
        if loop.is_running():
            # Can't run async in already running loop, use sync fallback
            session_path = get_session_path(user_id, test_id)
            if not session_path.exists():
                return None
            with open(session_path, 'r', encoding='utf-8') as f:
                return json.load(f)
        else:
            return loop.run_until_complete(load_test_session_safe(user_id, test_id))
    except Exception as e:
        log.error(f"Error in load_test_session wrapper: {e}")
        return None

def delete_test_session(user_id: int, test_id: str) -> bool:
    """Synchronous wrapper for delete_test_session_safe"""
    import asyncio
    try:
        loop = asyncio.get_event_loop()
        if loop.is_running():
            # Can't run async in already running loop, use sync fallback
            session_path = get_session_path(user_id, test_id)
            if session_path.exists():
                session_path.unlink()
                return True
            return False
        else:
            return loop.run_until_complete(delete_test_session_safe(user_id, test_id))
    except Exception as e:
        log.error(f"Error in delete_test_session wrapper: {e}")
        return False

def get_user_sessions(user_id: int) -> List[Dict]:
    """Wrapper for get_user_sessions_safe (already synchronous)"""
    return get_user_sessions_safe(user_id)

# ========== GROUP ADMIN MANAGEMENT ==========

async def sync_group_admins(group_id: int) -> Dict[int, dict]:
    """
    Sync Telegram group admins with local data.
    This fetches admins from Telegram and updates our local records.
    """
    try:
        from config import bot, OWNER_ID
        
        # Get current admins from Telegram
        chat_admins = await bot.get_chat_administrators(group_id)
        
        telegram_admins = {}
        for admin in chat_admins:
            if not admin.user.is_bot:
                telegram_admins[admin.user.id] = {
                    'username': admin.user.username or '',
                    'first_name': admin.user.first_name or '',
                    'last_name': admin.user.last_name or '',
                    'status': admin.status,
                    'is_creator': admin.status == 'creator'
                }
        
        # Load current admin data
        admins_data = load_admins()
        if 'group_admins' not in admins_data:
            admins_data['group_admins'] = {}
        
        # Process each Telegram admin
        for admin_id, admin_info in telegram_admins.items():
            admin_id_str = str(admin_id)
            
            # Skip bot owner (they're always super admin)
            if admin_id == OWNER_ID:
                continue
            
            # Update or create admin entry
            if admin_id_str not in admins_data['group_admins']:
                admins_data['group_admins'][admin_id_str] = {
                    'groups': [],
                    'name': f"{admin_info['first_name']} {admin_info['last_name']}".strip(),
                    'username': admin_info['username'],
                    'can_create': True,
                    'can_activate': True
                }
            
            # Add this group to admin's list if not present
            if group_id not in admins_data['group_admins'][admin_id_str]['groups']:
                admins_data['group_admins'][admin_id_str]['groups'].append(group_id)
            
            # Update admin info
            admins_data['group_admins'][admin_id_str]['username'] = admin_info['username']
            admins_data['group_admins'][admin_id_str]['name'] = f"{admin_info['first_name']} {admin_info['last_name']}".strip()
        
        # Remove admins who are no longer admins in Telegram
        for admin_id_str in list(admins_data['group_admins'].keys()):
            admin_data = admins_data['group_admins'][admin_id_str]
            if group_id in admin_data['groups']:
                if int(admin_id_str) not in telegram_admins:
                    admin_data['groups'].remove(group_id)
                    # If admin has no groups left, remove them
                    if not admin_data['groups']:
                        del admins_data['group_admins'][admin_id_str]
        
        # Save updated admin data
        save_admins(admins_data)
        
        log.info(f"Synced {len(telegram_admins)} admins for group {group_id}")
        return telegram_admins
        
    except Exception as e:
        log.error(f"Failed to sync admins for group {group_id}: {e}")
        return {}

def is_group_admin(user_id: int, group_id: int) -> bool:
    """Check if user is admin for a specific group"""
    if is_owner(user_id):
        return True
    
    admins_data = load_admins()
    group_admins = admins_data.get('group_admins', {})
    user_admin_data = group_admins.get(str(user_id), {})
    
    return group_id in user_admin_data.get('groups', [])

def get_user_admin_groups(user_id: int) -> List[int]:
    """Get all groups where user is admin"""
    if is_owner(user_id):
        return load_group_ids()  # Owner can access all groups
    
    admins_data = load_admins()
    group_admins = admins_data.get('group_admins', {})
    user_admin_data = group_admins.get(str(user_id), {})
    
    return user_admin_data.get('groups', [])

def can_user_create_test(user_id: int) -> bool:
    """Check if user can create tests"""
    if is_owner(user_id):
        return True
    
    admins_data = load_admins()
    group_admins = admins_data.get('group_admins', {})
    user_admin_data = group_admins.get(str(user_id), {})
    
    return user_admin_data.get('can_create', False) and len(user_admin_data.get('groups', [])) > 0

def can_user_manage_test(user_id: int, test_id: str) -> bool:
    """Check if user can manage a specific test"""
    if is_owner(user_id):
        return True
    
    test_data = read_test(test_id)
    if not test_data:
        return False
    
    test_groups = set()
    for g in test_data.get('groups', []):
        try:
            test_groups.add(int(g))
        except:
            pass
    
    user_admin_groups = set(get_user_admin_groups(user_id))
    
    # User can manage if they admin any group the test is assigned to
    return bool(test_groups.intersection(user_admin_groups))

def is_admin(user_id: int) -> bool:
    """Check if user is admin of any group"""
    if is_owner(user_id):
        return True
    
    admins_data = load_admins()
    group_admins = admins_data.get('group_admins', {})
    
    return str(user_id) in group_admins and len(group_admins[str(user_id)].get('groups', [])) > 0

async def get_bot_admin_groups(user_id: int) -> List[int]:
    """
    Get groups where the BOT is admin, filtered by user permissions.
    For owner: all groups where bot is admin
    For admins: only their assigned groups where bot is admin
    """
    try:
        from config import bot
        
        all_groups = load_group_ids()
        bot_admin_groups = []
        
        # Check each group to see if bot is admin
        for group_id in all_groups:
            try:
                bot_member = await bot.get_chat_member(group_id, (await bot.get_me()).id)
                if bot_member.status in ["administrator", "creator"]:
                    bot_admin_groups.append(group_id)
            except Exception as e:
                log.debug(f"Bot not admin in group {group_id}: {e}")
                continue
        
        # Filter based on user permissions
        if is_owner(user_id):
            # Owner can access all groups where bot is admin
            return bot_admin_groups
        else:
            # Admins can only access their assigned groups
            user_admin_groups = get_user_admin_groups(user_id)
            return [gid for gid in bot_admin_groups if gid in user_admin_groups]
    
    except Exception as e:
        log.error(f"Error getting bot admin groups: {e}")
        return []
# utils.py
def get_users_with_active_sessions(max_age_hours: int = 6) -> List[int]:
    """
    Get user IDs who have active test sessions.
    Only considers sessions updated within max_age_hours.
    """
    try:
        session_dir = Path("data/sessions")
        if not session_dir.exists():
            return []
        
        active_users = set()
        cutoff_time = time.time() - (max_age_hours * 3600)
        
        # Check session files
        for session_file in session_dir.glob("*.json"):
            try:
                # Check file modification time first
                if session_file.stat().st_mtime < cutoff_time:
                    continue
                
                # Parse filename to get user_id
                # Expecting format: session_<user_id>_<test_id>.json or <user_id>/<test_id>.json
                if session_file.name.startswith("session_"):
                    parts = session_file.stem.split("_")
                    if len(parts) >= 2:
                        user_id = int(parts[1])
                        active_users.add(user_id)
                elif session_file.parent.name.isdigit():
                    # User directory structure
                    user_id = int(session_file.parent.name)
                    active_users.add(user_id)
            except (ValueError, IndexError, OSError) as e:
                log.debug(f"Skipping session file {session_file}: {e}")
                continue
        
        # Also check user session directories
        for user_dir in session_dir.iterdir():
            if user_dir.is_dir() and user_dir.name.isdigit():
                try:
                    user_id = int(user_dir.name)
                    
                    # Check if user has any recent session files
                    for test_session in user_dir.glob("*.json"):
                        if test_session.stat().st_mtime >= cutoff_time:
                            active_users.add(user_id)
                            break
                except (ValueError, OSError):
                    continue
        
        return sorted(active_users)
        
    except Exception as e:
        log.error(f"Error getting users with active sessions: {e}")
        return []


async def notify_groups_and_members(group_ids: List[int], test_name: str, tid: str) -> Dict[str, int]:
    """
    Faol test yangi guruhlarga berilganda:
      1) Guruh chatiga umumiy e'lon
      2) Guruh a'zolariga shaxsiy DM

    Return:
      {
        "groups_notified": <int>,
        "total_notified": <int>,
        "total_failed": <int>
      }
    """
    # Normalize ids -> int
    groups = []
    for g in group_ids:
        try:
            groups.append(int(g))
        except Exception:
            continue

    if not groups:
        return {"groups_notified": 0, "total_notified": 0, "total_failed": 0}

    titles = load_group_titles()
    members_map = load_group_members()  # { "chat_id_str": { "members": [...], "member_data": {...} } }

    groups_notified = 0
    total_notified = 0
    total_failed = 0

    # Guruh e'lon matni
    def group_text(title: str) -> str:
        return (
            "🟢 <b>Test boshlandi</b>\n\n"
            f"📚 {test_name}\n"
            f"🆔 <code>{tid}</code>\n\n"
            "Imtihonni boshlash uchun: botga yozing yoki /start buyrug‘ini bosing."
        )

    # DM matni
    user_text = (
        f"📚 <b>{test_name}</b> testi guruhingizda boshlandi.\n"
        f"🆔 <code>{tid}</code>\n\n"
        "Boshlash: shu botga /start yuboring yoki menyudagi tugmalardan foydalaning."
    )

    for gid in groups:
        title = titles.get(gid, f"Guruh {gid}")

        # 1) Guruhga xabar
        try:
            await bot.send_message(gid, group_text(title), disable_web_page_preview=True)
            groups_notified += 1
        except Exception as e:
            log.warning(f"Group notify failed for {gid}: {e}")

        # 2) A'zolarga DM
        mems = members_map.get(str(gid), {}).get("members", []) or []
        # Agar sinxronlanmagan bo‘lsa, shunchaki DM qismi bo‘sh qoladi.
        # (xohlasang bu yerda fallback sifatida sync chaqirishingiz mumkin)

        for uid in mems:
            try:
                await bot.send_message(uid, user_text, disable_web_page_preview=True)
                total_notified += 1
            except Exception as e:
                total_failed += 1
                # Flood limitlarga tushmaslik uchun logni yumshoq qilamiz
                if total_failed <= 5:
                    log.debug(f"DM failed for {uid}: {e}")

        # Telegram floodni yumshatamiz
        await asyncio.sleep(0.2)

    return {
        "groups_notified": groups_notified,
        "total_notified": total_notified,
        "total_failed": total_failed,
    }

def get_test_active_groups(test_id: str) -> List[int]:
    obj = read_test(test_id) or {}
    out = []
    for g in obj.get("active_groups", []) or []:
        try:
            out.append(int(g))
        except Exception:
            pass
    return sorted(set(out))

def set_test_active_groups(test_id: str, groups: List[int]):
    obj = read_test(test_id) or {"test_id": test_id}
    norm = sorted(set(int(x) for x in groups))
    obj["active_groups"] = norm
    write_test(test_id, obj)


async def validate_and_sync_new_user(user_id: int) -> Tuple[bool, List[int]]:
    
    """
    Validate if a user is in any groups, syncing with owner account if needed.
    Specifically designed for users starting the bot for the first time.
    """
    try:
        # First check if user exists in our data
        user_groups = get_user_groups(user_id)
        gm = load_group_members()
        
        user_found_in_data = False
        for gid_str, rec in gm.items():
            if user_id in rec.get("members", []):
                user_found_in_data = True
                break
        
        if user_found_in_data or user_groups:
            # User exists, validate they're still in groups
            return await validate_user_still_in_groups(user_id)
        
        # User not in our data - they might be new
        log.info(f"User {user_id} not found in data, checking with owner account...")
        
        # Get owner's Telethon client
        from telethon_service import get_user_telethon_service
        user_telethon = await get_user_telethon_service()
        
        if not user_telethon:
            log.error("Owner Telethon not available for new user check")
            return False, []
        
        # Get all groups where bot is present
        groups = load_group_ids()
        found_groups = []
        
        for group_id in groups:
            try:
                # Fetch ALL members using owner account
                members, _ = await user_telethon.fetch_all_group_members(group_id)
                
                # Check if our user is in this group
                for member in members:
                    if member['id'] == user_id:
                        found_groups.append(group_id)
                        
                        # Add user to our data immediately
                        add_user_to_group(user_id, group_id)
                        
                        # Update member data
                        gm = load_group_members()
                        if str(group_id) not in gm:
                            gm[str(group_id)] = {"members": [], "member_data": {}}
                        
                        gm[str(group_id)]["member_data"][str(user_id)] = {
                            "first_name": member.get("first_name", ""),
                            "last_name": member.get("last_name", ""),
                            "username": member.get("username", ""),
                            "phone": member.get("phone", ""),
                            "is_premium": member.get("is_premium", False),
                            "is_bot": False,
                            "is_admin": False
                        }
                        write_json(Path(GROUP_MEMBERS_FILE), gm)
                        
                        log.info(f"Found and added new user {user_id} to group {group_id}")
                        break
                
                # Small delay to avoid rate limits
                await asyncio.sleep(0.2)
                
            except Exception as e:
                log.error(f"Error checking group {group_id} for user {user_id}: {e}")
                continue
        
        return len(found_groups) > 0, found_groups
        
    except Exception as e:
        log.error(f"Error in validate_and_sync_new_user for {user_id}: {e}")
        return False, []